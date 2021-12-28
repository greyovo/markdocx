# noinspection PyProtectedMember
#
import io
import os
from urllib.request import urlopen

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import *
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.styles import style
from docx.styles.style import _ParagraphStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run

debug_state: bool = False
auto_open: bool = True
show_image_desc: bool = True  # 是否显示图片的描述，即 `![desc](src/img)` 中 desc的内容

document: Document = Document()


def debug(*args):
    print(*args) if debug_state else None


def init_styles():
    styles = document.styles

    # 正文
    normal = styles['Normal']
    normal.font.name = "Calibri"  # 只设置name是设置西文字体
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 要额外设置中文字体

    # 各级标题
    # 直接对heading的样式设置是不生效的，需要新建一个同名样式覆盖
    for i in range(1, 5):
        new_style: _ParagraphStyle = styles.add_style('Heading%d' % i, WD_STYLE_TYPE.PARAGRAPH)
        new_style.base_style = styles['Heading %d' % i]
        new_style.quick_style = True
        new_style.font.name = "Times New Roman"  # 只设置name是设置西文字体
        new_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 要额外设置中文字体
        new_style.font.color.rgb = RGBColor(0, 0, 0)

    # 引用块
    blockquote: style = styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    # blockquote.base_style = styles["Normal"]
    blockquote.font.name = "Calibri"  # 只设置name是设置西文字体
    blockquote._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 要额外设置中文字体
    # block_format = blockquote.
    # blockquote.font.color.rgb = RGBColor(255, 0, 0)
    blockquote.font.italic = True


# h1, h2, ...
def add_heading(content: str, tag: str):
    level: int = int(tag.__getitem__(1))
    p = document.add_paragraph(content, style="Heading%d" % level)
    return p


# bold, italic, strike...
def add_run(p: Paragraph, content: str, char_style: str = "plain"):
    # fixme 行内的样式超过一个的句子会被忽略，如：
    # <u>**又加粗又*斜体*又下划线**</u>
    debug("[%s]:" % char_style, content)

    run = p.add_run(content)

    # 不应当使用形如 run.bold = (char_style=="strong") 的方式
    # 因为没有显式加粗，不意味着整体段落不加粗。
    if char_style == "strong":
        run.bold = True
    if char_style == "em":
        run.italic = True
    if char_style == "u":
        run.underline = True
    if char_style == "strike":
        run.font.strike = True
    if char_style == "sub":
        run.font.subscript = True
    if char_style == "sup":
        run.font.superscript = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW if char_style == "highlight" else None

    # TODO 代码块样式
    if char_style == "code":
        run.font.name = "Consolas"

    # TODO 引用块
    # TODO 有序列表、无序列表、TODO列表


def add_picture(elem):
    p: Paragraph = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run: Run = p.add_run()

    img_src: str

    if elem["src"] != "":
        img_src = elem["src"]
        debug("[local image]:", img_src)
        run.add_picture(img_src, width=Inches(5.8))
    else:
        img_src = elem["title"]
        # if img_src.startswith("http"):
        print("[fetching image...]:", img_src)
        image_bytes = urlopen(img_src).read()
        data_stream = io.BytesIO(image_bytes)
        run.add_picture(data_stream, width=Inches(5.8))

    # 如果选择展示图片描述，那么描述会在图片下方显示

    if show_image_desc and elem["alt"] != "":
        desc_p: Paragraph = document.add_paragraph(elem["alt"], style="Caption")  # TODO 图片描述的显示样式
        desc_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_table(table_root):
    # 统计列数
    col_count: int = 0
    for col in table_root.thead.tr.contents:
        if col.string != "\n":
            col_count += 1

    table = document.add_table(0, col_count, style="Table Grid")  # TODO 表格样式

    # 表格头行
    head_row_cells = table.add_row().cells
    i = 0
    for col in table_root.thead.tr.contents:
        if col.string == "\n":
            continue
        head_row_cells[i].paragraphs[0].add_run(col.string).bold = True
        # run = head_row_cells[i].add_paragraph(col.string)
        i += 1

    # 数据行
    for tr in table_root.tbody:
        if tr.string == "\n":
            continue
        row_cells = table.add_row().cells
        i = 0
        for td in tr.contents:
            if td.string == "\n":
                continue
            row_cells[i].text = td.string
            i += 1


def add_number_list(number_list):
    print(number_list.contents, "\n")
    # print(number_list.children, "\n")

    for item in number_list.children:
        if item.string == "\n":
            continue
        print(str(item.contents[0]))  # TODO 写入word中
        if item.ol is not None:  # 有子序列
            for item2 in item.ol.children:
                if item2.string == "\n":
                    continue
                print(" ", item2.string)  # TODO 写入word中


def add_bullet_list(bullet_list):
    # print(bullet_list.contents)
    for item in bullet_list.children:
        text: str = str(item.string)
        if text == "\n":
            continue
        # 有可能是TODO list
        if text.startswith("[ ]") or text.startswith("[x]"):
            add_todo_list(bullet_list)
            return
        print(item.contents[0])  # TODO 写入word中
        if hasattr(item, "ul") and item.ul is not None:  # 有子序列
            for item2 in item.ul.children:
                if item2.string == "\n":
                    continue
                print(" ", item2.string)  # TODO 写入word中


def add_todo_list(todo_list):
    for item in todo_list.children:
        if item.string == "\n":
            continue
        text: str = item.string
        if text.startswith("[x]"):
            print(text.replace("[x]", "[ √ ]", 1))  # TODO 写入word中
        if text.startswith("[ ]"):
            print(text.replace("[ ]", "[   ]", 1))  # TODO 写入word中


# TODO 分割线
def add_split_line(p: Paragraph):
    pass


# TODO 超链接
def add_link(p: Paragraph, text: str, src: str):
    debug("[link]:", text, "[src]:", src)
    run = p.add_run(text)


def add_paragraph(children, style: str = None):
    """
    :param style:
    :type children: list
    一个段落内的元素（包括图片）。根据有无样式来划分，组成一个列表。
    有样式文字如加粗、斜体、图片、等。
    如`I am plain _while_ he is **bold**`将转为：
    ["I am plain", "while", "he is", "bold"]

    """
    p = document.add_paragraph(style=style)
    for elem in children.contents:  # 遍历一个段落内的所有元素
        if elem.name == "a":
            add_link(p, elem.string, elem["href"])
        elif elem.name == "img":
            add_picture(elem)
        elif elem.name is not None:  # 有字符样式的子串
            add_run(p, elem.string, elem.name)
        elif not elem.string == "\n":  # 无字符样式的子串
            add_run(p, elem)


# from docx.enum.style import WD_STYLE
def add_blockquote(children):
    debug(children.contents)
    for p in children.contents:
        if p.string != "\n":
            add_paragraph(p, style="Block Quote")


class DocxProcessing:
    def __init__(self, html_path: str, docx_path: str):
        # 打开HTML
        with open(html_path, 'r', encoding="UTF-8") as html_file:
            html_str = html_file.read()
        soup = BeautifulSoup(html_str, 'html.parser')
        body_tag = soup.contents[2]
        para: Paragraph
        # 初始化样式
        init_styles()
        # 逐个写到word中
        for root in body_tag.children:
            if root.string != "\n":
                # debug("<%s>" % root.name)

                # 判断是普通段落还是标题
                if root.name == "p":
                    add_paragraph(root)

                if root.name == "blockquote":
                    add_blockquote(root)

                if root.name == "ol":  # TODO 数字列表
                    # add_number_list(root)
                    pass

                if root.name == "ul":  # TODO 无序列表 或 TODO_List
                    # add_bullet_list(root)
                    pass

                if root.name == "table":
                    # fixme 测试看不同的效果而已
                    # for i in range(1, 7):
                    add_table(root)
                    pass

                if root.name == "h1" or root.name == "h2" or \
                        root.name == "h3" or root.name == "h4" or root.name == "h5":
                    para = add_heading(root.string, root.name)

        document.save(docx_path)
        print("docx saved to:", docx_path)


if __name__ == '__main__':
    _ = DocxProcessing("example.html", "example.docx")
    if auto_open:
        os.startfile("example.docx")
