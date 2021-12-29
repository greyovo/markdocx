from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

# 一段文本中的某几部分，通过run对象进行分割、单独设置样式
# A plain paragraph having some **bold** and some *italic*.
p = document.add_paragraph()
p.add_run('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True


document.add_heading('Heading, level 1', level=1)
document.add_paragraph(
    """In addition to collecting a set of formatting properties, a style has five properties that specify its behavior. This behavior is relatively simple, basically amounting to when and where the style appears in the Word or LibreOffice UI.The key notion to understanding style behavior is the recommended list. In the style pane in Word, the user can select which list of styles they want to see. One of these is named Recommended and is known as the recommended list. All five behavior properties affect some aspect of the style’s appearance in this list and in the style gallery.""")
document.add_heading('Heading, level 2', level=2)
document.add_heading('Heading, level 3', level=3)
document.add_heading('Heading, level 4', level=4)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('windows11.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')
